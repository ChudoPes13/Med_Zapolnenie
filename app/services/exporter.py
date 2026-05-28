from __future__ import annotations

import json
from io import BytesIO

from docx import Document

from app.schemas import EMK, EvidenceOut, FindingOut


def export_json_package(emk: EMK, findings: list[FindingOut], evidence: list[EvidenceOut]) -> str:
    return json.dumps(
        {
            "format": "medjarvis.visit.v1",
            "form": "043/u" if emk.clinical_focus == "dental" else "025/u-draft",
            "live_write": False,
            "doctor_confirmation_required": True,
            "emk": emk.model_dump(),
            "quality": [item.model_dump() for item in findings],
            "evidence": [item.model_dump() for item in evidence],
        },
        ensure_ascii=False,
        indent=2,
    )


def export_html(emk: EMK, findings: list[FindingOut], evidence: list[EvidenceOut]) -> str:
    rows = "".join(
        f"<li><strong>{item.title}</strong>: {item.message} [{item.status}]</li>"
        for item in findings
    )
    ev_rows = "".join(
        f"<li>{item.kr_id}, {item.section}: {item.fragment}</li>" for item in evidence
    )
    focus_block = ""
    if emk.clinical_focus == "dental":
        focus_block = (
            f"<p>Зуб FDI: {emk.dental.tooth_fdi or ''}; "
            f"перкуссия: {emk.dental.percussion or ''}; "
            f"термопроба: {emk.dental.thermal_test or ''}; "
            f"ЭОД: {emk.dental.eod_mka or ''}</p>"
        )
    if emk.clinical_focus == "lower_limb":
        focus_block = (
            f"<p>Нижние конечности: сторона {emk.lower_limb.side or ''}; "
            f"область {emk.lower_limb.location or ''}; "
            f"пульс стопы {emk.lower_limb.dorsalis_pedis_pulse or ''}; "
            f"кожа {emk.lower_limb.skin_color or ''} / "
            f"{emk.lower_limb.skin_temperature or ''}</p>"
        )
    return f"""<!doctype html>
<html lang="ru">
<head><meta charset="utf-8"><title>МедЖарвис</title></head>
<body>
<h1>МедЖарвис: прием</h1>
<h2>Жалобы</h2><p>{"; ".join(emk.complaints)}</p>
<h2>Анамнез</h2><p>{"; ".join(emk.anamnesis)}</p>
<h2>Объективно</h2><p>{"; ".join(emk.objective)}</p>
<h2>Профильный осмотр</h2>{focus_block}
<h2>Диагноз</h2><p>{emk.diagnosis.code or ""} {emk.diagnosis.title or ""}</p>
<h2>Итог</h2><p>{emk.final_summary or ""}</p>
<h2>Контроль качества</h2><ul>{rows}</ul>
<h2>КР Минздрава</h2><ul>{ev_rows}</ul>
</body></html>"""


def export_1c_text(emk: EMK, findings: list[FindingOut], evidence: list[EvidenceOut]) -> str:
    open_findings = [item.title for item in findings if item.status == "open"]
    return "\n".join(
        [
            "МЕДЖАРВИС 043/У",
            f"Жалобы: {'; '.join(emk.complaints)}",
            f"Анамнез: {'; '.join(emk.anamnesis)}",
            f"Объективно: {'; '.join(emk.objective)}",
            (
                "Профильный осмотр: "
                + (
                    f"зуб {emk.dental.tooth_fdi or '-'}, "
                    f"перкуссия {emk.dental.percussion or '-'}, "
                    f"ЭОД {emk.dental.eod_mka or '-'}"
                    if emk.clinical_focus == "dental"
                    else (
                        f"нижние конечности, сторона {emk.lower_limb.side or '-'}, "
                        f"область {emk.lower_limb.location or '-'}, "
                        f"пульс {emk.lower_limb.dorsalis_pedis_pulse or '-'}"
                    )
                )
            ),
            f"Диагноз: {emk.diagnosis.code or '-'} {emk.diagnosis.title or ''}",
            f"Итог: {emk.final_summary or '-'}",
            f"Рекомендации: {'; '.join(emk.recommendations)}",
            f"Открытые замечания: {'; '.join(open_findings) if open_findings else 'нет'}",
            f"Источники КР: {'; '.join(item.kr_id for item in evidence)}",
            "Live-write: выключен до подтверждения врача",
        ]
    )


def export_docx_bytes(emk: EMK, findings: list[FindingOut], evidence: list[EvidenceOut]) -> bytes:
    document = Document()
    document.add_heading("МедЖарвис: прием", 0)
    document.add_heading("Жалобы", level=1)
    document.add_paragraph("; ".join(emk.complaints) or "-")
    document.add_heading("Объективно", level=1)
    document.add_paragraph("; ".join(emk.objective) or "-")
    if emk.clinical_focus == "dental":
        document.add_paragraph(f"Зуб FDI: {emk.dental.tooth_fdi or '-'}")
        document.add_paragraph(f"Перкуссия: {emk.dental.percussion or '-'}")
        document.add_paragraph(f"Термопроба: {emk.dental.thermal_test or '-'}")
        document.add_paragraph(f"ЭОД: {emk.dental.eod_mka or '-'}")
    if emk.clinical_focus == "lower_limb":
        document.add_paragraph(f"Сторона: {emk.lower_limb.side or '-'}")
        document.add_paragraph(f"Область: {emk.lower_limb.location or '-'}")
        document.add_paragraph(f"Пульс стопы: {emk.lower_limb.dorsalis_pedis_pulse or '-'}")
        document.add_paragraph(
            f"Кожа: {emk.lower_limb.skin_color or '-'} / {emk.lower_limb.skin_temperature or '-'}"
        )
    document.add_heading("Диагноз", level=1)
    document.add_paragraph(f"{emk.diagnosis.code or '-'} {emk.diagnosis.title or ''}")
    document.add_heading("Итог", level=1)
    document.add_paragraph(emk.final_summary or "-")
    document.add_heading("Контроль качества", level=1)
    for item in findings:
        document.add_paragraph(f"{item.title}: {item.message} [{item.status}]", style="List Bullet")
    document.add_heading("КР Минздрава", level=1)
    for item in evidence:
        document.add_paragraph(
            f"{item.kr_id} / {item.section}: {item.fragment}",
            style="List Bullet",
        )

    output = BytesIO()
    document.save(output)
    return output.getvalue()
